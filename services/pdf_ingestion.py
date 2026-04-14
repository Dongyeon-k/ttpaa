from __future__ import annotations

import logging
import re
from pathlib import Path

from django.conf import settings
from django.db import transaction
from django.utils import timezone
from pypdf import PdfReader

from apps.chatbot.models import ConstitutionChunk, ConstitutionDocument, ConstitutionPage

logger = logging.getLogger("ttpaa")

HEADING_PATTERN = re.compile(r"(제\s*\d+\s*[장조]|제\s*\d+\s*[장]|제\s*\d+\s*조|Chapter\s+\d+|Article\s+\d+)")
ALLOWED_CONSTITUTION_EXTENSIONS = {".pdf", ".docx", ".xlsx"}


def get_constitution_file_extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def _ocr_pdf_page(pdf_bytes: bytes, page_number: int) -> str:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("OCR Python dependencies are not installed. Run pip install -r requirements.txt.") from exc

    dpi = getattr(settings, "CONSTITUTION_OCR_DPI", 200)
    lang = getattr(settings, "CONSTITUTION_OCR_LANG", "kor+eng")
    zoom = dpi / 72

    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf:
        page = pdf.load_page(page_number - 1)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
        return (pytesseract.image_to_string(image, lang=lang) or "").strip()


def _split_chunks(text: str, chunk_size: int = 900) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n{2,}", text) if part.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs or [text]:
        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current:
                chunks.append(current)
            current = paragraph
    if current:
        chunks.append(current)
    return chunks or [text]


def _extract_heading(text: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if HEADING_PATTERN.search(line):
            return line[:255]
    return ""


def _extract_pdf_pages(document: ConstitutionDocument) -> list[tuple[int, str]]:
    document.file.open("rb")
    reader = PdfReader(document.file)
    pdf_bytes = None
    extracted_pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        min_text_chars = getattr(settings, "CONSTITUTION_OCR_MIN_TEXT_CHARS", 20)
        if len(text) < min_text_chars and getattr(settings, "CONSTITUTION_OCR_ENABLED", True):
            if pdf_bytes is None:
                document.file.seek(0)
                pdf_bytes = document.file.read()
            try:
                ocr_text = _ocr_pdf_page(pdf_bytes, page_number)
            except Exception:
                if text:
                    logger.warning("constitution_ocr_fallback_failed document_id=%s page=%s", document.pk, page_number)
                    ocr_text = ""
                else:
                    raise
            if len(ocr_text) > len(text):
                text = ocr_text
        extracted_pages.append((page_number, text))
    return extracted_pages


def _extract_docx_pages(document: ConstitutionDocument) -> list[tuple[int, str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX ingestion dependency is not installed. Run pip install -r requirements.txt.") from exc

    document.file.open("rb")
    docx_document = Document(document.file)
    parts = [paragraph.text.strip() for paragraph in docx_document.paragraphs if paragraph.text.strip()]
    for table in docx_document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return [(1, "\n".join(parts).strip())]


def _format_spreadsheet_value(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _extract_xlsx_pages(document: ConstitutionDocument) -> list[tuple[int, str]]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("XLSX ingestion dependency is not installed. Run pip install -r requirements.txt.") from exc

    document.file.open("rb")
    workbook = load_workbook(document.file, read_only=True, data_only=True)
    try:
        extracted_pages = []
        for page_number, worksheet in enumerate(workbook.worksheets, start=1):
            rows = []
            for row in worksheet.iter_rows(values_only=True):
                values = [_format_spreadsheet_value(value) for value in row]
                while values and not values[-1]:
                    values.pop()
                if any(values):
                    rows.append("\t".join(values))
            sheet_text = "\n".join(rows).strip()
            if sheet_text:
                sheet_text = f"[시트: {worksheet.title}]\n{sheet_text}"
            extracted_pages.append((page_number, sheet_text))
        return extracted_pages
    finally:
        workbook.close()


def _extract_document_pages(document: ConstitutionDocument) -> tuple[list[tuple[int, str]], str]:
    extension = get_constitution_file_extension(document.upload_filename or document.file.name)
    if extension == ".pdf":
        return _extract_pdf_pages(document), "PDF"
    if extension == ".docx":
        return _extract_docx_pages(document), "DOCX"
    if extension == ".xlsx":
        return _extract_xlsx_pages(document), "XLSX"
    raise ValueError("PDF, DOCX, XLSX 파일만 인덱싱할 수 있습니다.")


def index_constitution_document(document: ConstitutionDocument) -> ConstitutionDocument:
    logger.info("constitution_index_started document_id=%s", document.pk)
    document.index_status = ConstitutionDocument.STATUS_PENDING
    document.index_error = ""
    document.save(update_fields=["index_status", "index_error"])

    try:
        extracted_pages, file_type_label = _extract_document_pages(document)
        total_text_length = sum(len(text) for _, text in extracted_pages)
        if not extracted_pages or total_text_length == 0:
            raise ValueError(
                f"{file_type_label}에서 검색 가능한 텍스트를 추출하지 못했습니다. 파일 내용을 확인해 주세요."
            )

        with transaction.atomic():
            document.pages.all().delete()
            document.chunks.all().delete()
            for page_number, text in extracted_pages:
                page_obj = ConstitutionPage.objects.create(document=document, page_number=page_number, text=text)
                if not text:
                    continue
                heading = _extract_heading(text)
                for chunk_index, chunk_text in enumerate(_split_chunks(text), start=1):
                    if not chunk_text.strip():
                        continue
                    ConstitutionChunk.objects.create(
                        document=document,
                        page=page_obj,
                        page_number=page_number,
                        chunk_index=chunk_index,
                        heading=heading,
                        text=chunk_text,
                    )
        document.index_status = ConstitutionDocument.STATUS_INDEXED
        document.indexed_at = timezone.now()
        document.index_error = ""
        document.save(update_fields=["index_status", "indexed_at", "index_error"])
        logger.info("constitution_index_completed document_id=%s", document.pk)
        return document
    except Exception as exc:
        document.index_status = ConstitutionDocument.STATUS_FAILED
        document.index_error = str(exc)
        document.save(update_fields=["index_status", "index_error"])
        logger.exception("constitution_index_failed document_id=%s", document.pk)
        raise
